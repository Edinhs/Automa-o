import re
import os
from typing import Dict, Any, Tuple
# Observação: python-docx pode não estar no venv padrão na fase de testes iniciais,
# mas o import é seguro se executado no venv correto. Tratamos exceções de import.
try:
    import docx
except ImportError:
    docx = None

# Regex para detectar o padrão VF_V_R (ex: VF395_V1_R1 ou VF429_V4_R1)
VF_PATTERN = re.compile(r"VF\d+_V\d+_R\d+")
VF_PARSED_PATTERN = re.compile(r"VF(\d+)_V(\d+)_R(\d+)")

def parse_vf_name(vf_name: str) -> Tuple[str, str]:
    """Extrai (versão, revisão) de um nome de VF como VF395_V1_R1."""
    match = VF_PARSED_PATTERN.match(vf_name)
    if match:
        return f"V{match.group(2)}", f"R{match.group(3)}"
    return "V1", "R1"

def extract_vfs_from_word(word_file_path: str) -> Dict[str, Dict[str, Any]]:
    """
    Lê o documento Word de VFs, localiza as VFs no INDEX e extrai seu conteúdo.
    Retorna um dicionário:
    {
       "VF395_V1_R1": {
          "content": "Conteúdo de texto da VF...",
          "version": "V1",
          "revision": "R1"
       }
    }
    """
    if not docx:
        raise ImportError("A biblioteca 'python-docx' é necessária para ler documentos Word. Instale-a via requirements.txt.")

    if not os.path.exists(word_file_path):
        raise FileNotFoundError(f"Arquivo Word não encontrado: {word_file_path}")

    doc = docx.Document(word_file_path)
    
    # 1. Varre o documento para identificar todas as VFs mencionadas e seu conteúdo
    # Como as VFs são seções do documento, vamos agrupar os parágrafos correspondentes a cada VF.
    vfs_data: Dict[str, list] = {}
    current_vf = None

    # Varre todos os parágrafos do documento
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Busca por ocorrência da nomenclatura da VF no início do parágrafo ou como cabeçalho
        matches = VF_PATTERN.findall(text)
        if matches:
            # Pegamos a primeira correspondência que parece ser um título de seção de VF
            # O texto do parágrafo geralmente começa com a VF ou a contém em destaque
            for match in matches:
                # Se for o início de uma nova seção de VF
                current_vf = match
                if current_vf not in vfs_data:
                    vfs_data[current_vf] = []
                # Evita duplicar a linha de título se já processamos
                vfs_data[current_vf].append(text)
                break
        elif current_vf:
            vfs_data[current_vf].append(text)

    # Varre também tabelas se houver, buscando associá-las a VFs
    # (Opcional, mas útil caso o conteúdo da VF esteja em tabelas)
    # Procuramos o contexto da VF ativa antes da tabela
    for table in doc.tables:
        table_text_lines = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip())
            if row_text:
                table_text_lines.append(row_text)
        
        # Se encontramos uma tabela e temos uma VF ativa, anexamos a tabela a ela
        if current_vf and table_text_lines:
            vfs_data[current_vf].append("\n".join(table_text_lines))

    # Formata o resultado final
    processed_vfs: Dict[str, Dict[str, Any]] = {}
    for vf_name, text_lines in vfs_data.items():
        content = "\n".join(text_lines)
        version, revision = parse_vf_name(vf_name)
        processed_vfs[vf_name] = {
            "content": content,
            "version": version,
            "revision": revision
        }

    # Caso a extração por parágrafos falhe em mapear as VFs devido a formatação específica do INDEX,
    # podemos fazer um fallback lendo especificamente as tabelas do INDEX
    if not processed_vfs:
        # Tenta varrer as primeiras tabelas do documento que costumam conter o INDEX
        for table in doc.tables[:3]: # primeiras tabelas geralmente são do sumário
            for row in table.rows:
                for cell in row.cells:
                    matches = VF_PATTERN.findall(cell.text)
                    for match in matches:
                        if match not in processed_vfs:
                            version, revision = parse_vf_name(match)
                            processed_vfs[match] = {
                                "content": f"VF extraída do INDEX: {match}",
                                "version": version,
                                "revision": revision
                            }

    return processed_vfs
